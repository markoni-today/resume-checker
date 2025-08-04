"""
Парсер файлов для Resume Checker
Поддерживает TXT, PDF, DOCX, DOC форматы
"""

import io
import re
from typing import Optional, Dict, Any
from pathlib import Path

try:
    import PyPDF2
    from docx import Document
except ImportError:
    PyPDF2 = None
    Document = None


class FileParser:
    """Парсер различных форматов файлов резюме"""
    
    SUPPORTED_TYPES = {
        'text/plain': 'txt',
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc'
    }
    
    def __init__(self):
        self.max_file_size = 5 * 1024 * 1024  # 5MB
    
    def parse_file(self, file_content: bytes, content_type: str, filename: str) -> str:
        """
        Парсит файл и возвращает текстовое содержимое
        
        Args:
            file_content: Содержимое файла в байтах
            content_type: MIME тип файла
            filename: Имя файла
            
        Returns:
            str: Извлеченный текст
            
        Raises:
            ValueError: Если файл не поддерживается или поврежден
        """
        
        # Проверка размера файла
        if len(file_content) > self.max_file_size:
            raise ValueError(f"Файл слишком большой. Максимальный размер: {self.max_file_size // 1024 // 1024}MB")
        
        # Определение типа файла
        file_type = self._detect_file_type(content_type, filename)
        
        # Парсинг в зависимости от типа
        try:
            if file_type == 'txt':
                return self._parse_txt(file_content)
            elif file_type == 'pdf':
                return self._parse_pdf(file_content)
            elif file_type == 'docx':
                return self._parse_docx(file_content)
            elif file_type == 'doc':
                return self._parse_doc(file_content)
            else:
                raise ValueError(f"Неподдерживаемый тип файла: {file_type}")
                
        except Exception as e:
            raise ValueError(f"Ошибка при парсинге файла: {str(e)}")
    
    def _detect_file_type(self, content_type: str, filename: str) -> str:
        """Определяет тип файла по MIME типу и расширению"""
        
        # Проверка по MIME типу
        if content_type in self.SUPPORTED_TYPES:
            return self.SUPPORTED_TYPES[content_type]
        
        # Проверка по расширению файла
        extension = Path(filename).suffix.lower()
        extension_map = {
            '.txt': 'txt',
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc'
        }
        
        if extension in extension_map:
            return extension_map[extension]
        
        raise ValueError(f"Неподдерживаемый формат файла: {extension}")
    
    def _parse_txt(self, file_content: bytes) -> str:
        """Парсит TXT файл"""
        try:
            # Попытка декодирования в UTF-8
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Попытка декодирования в Windows-1251 (популярная кодировка в России)
                return file_content.decode('windows-1251')
            except UnicodeDecodeError:
                try:
                    # Попытка декодирования в cp1252
                    return file_content.decode('cp1252')
                except UnicodeDecodeError:
                    raise ValueError("Не удалось определить кодировку текстового файла")
    
    def _parse_pdf(self, file_content: bytes) -> str:
        """Парсит PDF файл"""
        if PyPDF2 is None:
            raise ValueError("PyPDF2 не установлен. Невозможно обработать PDF файлы.")
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise ValueError("PDF файл не содержит страниц")
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            
            if not text_parts:
                raise ValueError("Не удалось извлечь текст из PDF файла")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            if "PDF" in str(e):
                raise ValueError(f"Ошибка при чтении PDF: {str(e)}")
            else:
                raise ValueError("Поврежденный или защищенный PDF файл")
    
    def _parse_docx(self, file_content: bytes) -> str:
        """Парсит DOCX файл"""
        if Document is None:
            raise ValueError("python-docx не установлен. Невозможно обработать DOCX файлы.")
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Извлечение текста из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            if not text_parts:
                raise ValueError("DOCX файл не содержит текста")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            if "docx" in str(e).lower():
                raise ValueError(f"Ошибка при чтении DOCX: {str(e)}")
            else:
                raise ValueError("Поврежденный DOCX файл")
    
    def _parse_doc(self, file_content: bytes) -> str:
        """Парсит DOC файл (старый формат Microsoft Word)"""
        # Для простоты MVP, DOC файлы не поддерживаются полностью
        # В production можно добавить поддержку через python-docx2txt или antiword
        raise ValueError("DOC формат пока не поддерживается. Пожалуйста, конвертируйте файл в DOCX.")
    
    def extract_metadata(self, file_content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
        """Извлекает метаданные файла"""
        metadata = {
            'filename': filename,
            'content_type': content_type,
            'size_bytes': len(file_content),
            'size_formatted': self._format_file_size(len(file_content))
        }
        
        file_type = self._detect_file_type(content_type, filename)
        metadata['file_type'] = file_type
        
        return metadata
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Форматирует размер файла в человекочитаемый вид"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"


def validate_file_for_parsing(file_content: bytes, content_type: str, filename: str) -> Dict[str, Any]:
    """
    Валидирует файл перед парсингом
    
    Returns:
        Dict с результатами валидации
    """
    parser = FileParser()
    
    try:
        # Проверка размера
        if len(file_content) > parser.max_file_size:
            return {
                'valid': False,
                'error': f'Файл слишком большой. Максимальный размер: {parser.max_file_size // 1024 // 1024}MB'
            }
        
        # Проверка типа файла
        file_type = parser._detect_file_type(content_type, filename)
        
        # Проверка на пустой файл
        if len(file_content) == 0:
            return {
                'valid': False,
                'error': 'Файл пустой'
            }
        
        return {
            'valid': True,
            'file_type': file_type,
            'size': len(file_content)
        }
        
    except ValueError as e:
        return {
            'valid': False,
            'error': str(e)
        }